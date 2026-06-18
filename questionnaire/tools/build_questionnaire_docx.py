#!/usr/bin/env python3
"""Generate a branded, customer-facing Word version of the Snowflake sizing questionnaire.

Source of truth: the QUESTIONS data below (kept in sync with build_questionnaire_xlsx.py).
Output:          questionnaire/snowflake-sizing-questionnaire.docx

Word has no live data-validation/dropdowns, so enum answers are captured as a
shaded fill-in cell with the allowed options shown as guidance text.

Re-run after editing the questionnaire to regenerate the document:
    python3 questionnaire/tools/build_questionnaire_docx.py
"""
from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Paths
# --------------------------------------------------------------------------- #
HERE = os.path.dirname(os.path.abspath(__file__))
QUESTIONNAIRE_ROOT = os.path.abspath(os.path.join(HERE, ".."))
PLUGIN_ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
REPO_ROOT = os.path.abspath(os.path.join(PLUGIN_ROOT, "..", ".."))
CHANGELOG_PATH = os.path.join(PLUGIN_ROOT, "CHANGELOG.md")
LOGO_PATH = os.path.join(
    REPO_ROOT,
    "assets",
    "snowflake-branding",
    "snowflake-icons",
    "graphic_snowflake_logo_blue.png",
)
OUTPUT_PATH = os.path.join(
    QUESTIONNAIRE_ROOT, "snowflake-sizing-questionnaire.docx"
)

# --------------------------------------------------------------------------- #
# Brand palette (assets/snowflake-branding/snowflake.com/palette.json)
# --------------------------------------------------------------------------- #
BLUE = "29B5E8"        # primary
DARK_BLUE = "11567F"   # header
GREY = "8A999E"
ORANGE = "FF9F36"      # accent
WHITE = "FFFFFF"
ANSWER_FILL = "EAF7FC"   # very light blue - the fill-in column
ROW_TINT = "F4FBFE"      # alternating row tint
BORDER_GREY = "D5DDE0"

FONT = "Calibri"
BODY_PT = 14

# --------------------------------------------------------------------------- #
# Validation answer-type tokens (kept in sync with build_questionnaire_xlsx.py)
# --------------------------------------------------------------------------- #
CLOUD = ["AWS", "Azure", "Google Cloud"]
EDITION = ["Standard", "Enterprise", "Business Critical", "VPS"]
YESNO = ["Yes", "No"]
COMPLEXITY = ["Simple lookups", "Medium (joins + aggregations)", "Highly complex analytical"]
INGEST_FREQ = ["Batch", "Micro-batch", "Stream"]

# answer_type values:
#   "text"            -> free text, informational hint only
#   "num"             -> number >= 0
#   "pct"             -> number 0..100
#   ("list", [...])   -> options shown as guidance
#   ("list+", [...])  -> options shown as guidance (other values allowed)

# (id, category, question, reason, answer_type, unit_hint)
QUESTIONS = [
    # 1 - General & Strategic
    ("1a", "General & Strategic",
     "What are the primary business objectives for this data initiative? (e.g., reduce reporting time, enable new analytics, consolidate platforms)",
     "Helps understand the project's priority and informs architectural recommendations and potential high-value use cases.",
     "text", ""),
    ("1b", "General & Strategic",
     "Are you migrating from an existing data platform (e.g., Teradata, Oracle, Hadoop, Redshift)? If so, please specify.",
     "Informs the migration strategy and complexity, and allows more accurate estimates if the current platform's resource consumption data is available.",
     "text", ""),
    ("1c", "General & Strategic",
     'Which Cloud Provider (AWS / Azure / Google Cloud) and Region do you require? (Answer as: cloud + full region, e.g. "AWS Europe (London)")',
     "Determines the per-credit, AI-credit, and storage rates and ensures data-residency requirements are met. A fundamental component of every cost calculation.",
     ("list+", CLOUD), "cloud + full region"),
    ("1d", "General & Strategic",
     "Which Snowflake Edition do you require? (Standard / Enterprise / Business Critical / VPS)",
     "Sets the credit rate and the floor of available features. If unsure, leave blank and we infer it.",
     ("list", EDITION), ""),
    ("1e", "General & Strategic",
     "What is the proposed timeline? (development start [month/date], production go-live [month/date], and any phased rollout milestones)",
     "The go-live date and ramp shape drive the per-month consumption curve. A concrete dev-start and go-live month make the Year-1 ramp accurate instead of assumed.",
     "text", "dev start, go-live, milestones"),

    # 2 - Data Sources & Ingestion
    ("2a", "Data Sources & Ingestion",
     "What are the ongoing source systems for data? (e.g., on-prem databases, SaaS apps like Salesforce, IoT event streams, logs)",
     "Identifies the variety and complexity of data to ingest, informing the required ingestion tools and patterns (ETL/ELT).",
     "text", ""),
    ("2b", "Data Sources & Ingestion",
     "What are the formats of the incoming data? (structured CSV/Parquet, semi-structured JSON/XML, unstructured PDF/images). If known, give expected compression ratio (default 3x).",
     "Influences storage requirements and the compute needed for parsing/transformation. Format drives the compression ratio that converts raw TB to billed storage TB.",
     "text", "formats; compression ratio"),
    ("2c", "Data Sources & Ingestion",
     "For each of your largest data sources, what is the ingestion volume and frequency? ([GB/day] or [GB/month] + batch / micro-batch / stream)",
     "Directly impacts ingestion compute sizing. Volume drives ingestion credits and Snowpipe charges; frequency decides batch warehouse vs Snowpipe vs streaming.",
     "text", "[GB/day or GB/month] + frequency"),
    ("2d", "Data Sources & Ingestion",
     "Do any sources need to arrive within seconds or a few minutes of being generated? (live fraud alerts, real-time dashboards, IoT), if yes, expected [streaming GB/month].",
     "Determines whether Snowpipe Streaming (sub-second) is required vs standard Snowpipe. Streaming adds ingestion compute and the streaming GB/month sizes it.",
     ("list+", YESNO), "if yes: [streaming GB/month]"),

    # 3, Storage & Data Architecture
    ("3a", "Storage & Data Architecture",
     "What is the total raw source data volume [GB or TB] you plan to migrate initially? (if known, give expected compression ratio; default 3x)",
     "The primary driver of baseline storage cost. Compression reduces it, but raw size is the starting point for the estimate.",
     "text", "[GB or TB]; compression ratio"),
    ("3b", "Storage & Data Architecture",
     "How many [days] of historical data must be immediately available for query rollback (Time Travel)?",
     "Determines Time Travel and Fail-safe storage cost. Longer retention increases storage consumption and cost.",
     "num", "[days]"),
    ("3c", "Storage & Data Architecture",
     "Do you require separate, isolated environments for Development, Testing/QA, and Production? (which environments)",
     "Determines the number of databases/objects to manage and informs the compute strategy for multiple teams and release cycles.",
     "text", "which environments"),
    ("3d", "Storage & Data Architecture",
     "Do you have large volumes of historical data rarely or never queried after a certain age?, if yes: [% of total storage] and after how many [months] it becomes cold.",
     "Archive storage is significantly cheaper than standard. If cold data is a large proportion, enabling archive storage meaningfully reduces long-term TCO.",
     "text", "if yes: [% of storage] + [months]"),
    ("3e", "Storage & Data Architecture",
     "What proportion of your data changes (insert/update/delete) on a typical day? [% per day] (default 10%)",
     "Churn rate drives the Time Travel and Fail-safe overhead on top of base storage. Without it, overhead is assumed at 10% and may be materially wrong.",
     "pct", "[% per day]"),

    # 4, Workloads & Use Cases
    ("4a", "Workloads & Use Cases",
     "What are the operating hours for your environment overall? (weekdays 9am-6pm; 24x7; overnight batch only; mixed)",
     "A significant cost driver and the single most impactful missing input. Weekday-only = 22 days/month; 24x7 = 30 days/month.",
     "text", ""),
    ("4b", "Workloads & Use Cases",
     "List your primary workloads (Corporate BI, Ad-hoc Analytics, ELT/Transformation, Data Science), then answer 4c-4h for each one.",
     "Critical input for compute estimation. Separating workloads lets us right-size individual warehouses instead of one oversized warehouse.",
     "text", ""),
    ("4c", "Workloads & Use Cases",
     "[For each workload] How many users run queries concurrently during peak hours? [# concurrent users]",
     "Determines warehouse size and multi-cluster configuration needed to meet concurrency without queueing",
     "text", "[# concurrent users]"),
    ("4d", "Workloads & Use Cases",
     '[For each workload] What is the query-performance SLA? ([seconds] target, e.g. "dashboards < 5s", "ETL < 1hr")',
     "Drives warehouse size. Faster SLAs require larger warehouses.",
     "text", "[seconds] target"),
    ("4e", "Workloads & Use Cases",
     "[For each workload] What are the operating hours / frequency? (24x7 loading, 9am-5pm Mon-Fri BI, once daily 2am batch)",
     "Directly impacts credit consumption. Warehouses suspend when idle, so the schedule (hours/day x days/month) is key to cost.",
     "text", ""),
    ("4f", "Workloads & Use Cases",
     "[For each workload] How would you classify query complexity? (simple lookups / medium / highly complex analytical)",
     "Informs warehouse size. More complex queries need more compute (larger warehouses) to finish within the SLA.",
     ("list+", COMPLEXITY), "per workload"),
    ("4h", "Workloads & Use Cases",
     "[For each workload] Roughly how many queries per user per day [#], and typical query runtime [seconds] if known?",
     "Completes the BI/analytics credit formula (users x queries/day x runtime x size-credits x days). Without it, per-user consumption is assumed from complexity alone.",
     "text", "[# queries/user/day] + [seconds]"),

    # 5, Python, ML & Containers
    ("5a", "Python, Machine Learning & Containers",
     "Will users write and run Python directly inside the platform? (notebooks, Pandas/NumPy, scikit-learn, PySpark-style transforms)",
     "Python notebooks on standard warehouses suffice for exploration. Heavy in-memory ML benefits from Snowpark-Optimized Warehouses. Determines standard vs Snowpark-Optimized sizing.",
     ("list+", YESNO), ""),
    ("5b", "Python, Machine Learning & Containers",
     "Will you train ML models where large datasets load entirely into memory?, if yes: dataset size [GB] and training frequency.",
     "ML training at scale needs Snowpark-Optimized Warehouses. Dataset size and frequency prevent under-sizing.",
     "text", "if yes: [GB] + frequency"),
    ("5c", "Python, Machine Learning & Containers",
     "Do you need to run containerized apps/services inside the platform? (real-time inference API, web service, Docker app), if yes: app type, [requests/day], GPU or CPU.",
     "Containerized services run on Snowpark Container Services compute pools, separate from warehouses. App type, request volume, and GPU-vs-CPU are required to size the pool.",
     "text", "if yes: app type, [requests/day], GPU/CPU"),

    # 6, AI & Analytics Intelligence
    ("6a", "AI & Analytics Intelligence",
     "Would business users benefit from asking questions about data in plain English? (Cortex Analyst), if yes: [# analyst users] and [questions per user per day].",
     "Determines Cortex Analyst usage. User count x daily question frequency is the sizing input.",
     "text", "if yes: [# users] + [questions/user/day]"),
    ("6b", "AI & Analytics Intelligence",
     "Will you use an LLM to process or generate text in pipelines/apps? (summarising feedback, generating descriptions, classifying emails), if yes: [records/month] and avg text length.",
     "Determines Cortex Complete token volume. Cost = tokens x model rate.",
     "text", "if yes: [records/month] + text length"),
    ("6c", "AI & Analytics Intelligence",
     "Do you need to search across large collections of free-text documents by meaning? (contracts, tickets, KB articles), if yes: [indexed data GB] and refresh frequency.",
     "Determines Cortex Search, billed at X credits per GB of indexed data per month. The indexed GB is the key sizing input.",
     "text", "if yes: [indexed GB] + refresh freq"),
    ("6d", "AI & Analytics Intelligence",
     "Do you need to extract structured information from documents, PDFs, or images?, if yes: [documents/month] and [avg pages per document].",
     "Determines AI_EXTRACT token volume, billed per token from the AI credit pool. Document count x page length is required to estimate it.",
     "text", "if yes: [docs/month] + [pages/doc]"),
    ("6e", "AI & Analytics Intelligence",
     "Do you need text analysis on columns of data? (sentiment, classification, translation, summarisation), if yes: which functions, [# rows], avg text length.",
     "Determines which AI SQL functions are used (ai_sentiment / ai_classify / ai_translate / ai_summarize), each billed per million tokens.",
     "text", "if yes: functions, [# rows], text length"),
    ("6f", "AI & Analytics Intelligence",
     "Do you plan a conversational AI assistant / chatbot that answers by querying your data? (Cortex Agents / Snowflake Intelligence), if yes: [# users], [sessions/user/day], [messages/session].",
     "Determines Cortex Agents usage, billed per token. Usage is modeled as users x sessions/day x messages/session.",
     "text", "if yes: [# users], [sessions/day], [msgs/session]"),
    ("6g", "AI & Analytics Intelligence",
     "Do developers write SQL/Python regularly, and would an in-tool AI assistant be valuable? (Cortex Code), if yes: [# developers], surface(s) CLI / Snowsight / Desktop, [queries/dev/day].",
     "Determines Cortex Code, billed per token. Developer count, surface, and daily query volume are the key inputs.",
     "text", "if yes: [# devs], surface, [queries/dev/day]"),

    # 7, Data Governance & Security
    ("7a", "Data Governance & Security",
     "Do you have specific regulatory/compliance requirements? (HIPAA, PCI-DSS, GDPR)",
     "Determines the required Edition. Business Critical is necessary for the highest compliance levels (HIPAA, PCI).",
     "text", ""),
    ("7b", "Data Governance & Security",
     "Do you require dynamic data masking, row-level access policies, or tokenization?",
     "These are Enterprise-Edition-or-higher features. The answer sets the minimum required edition and its cost.",
     ("list+", YESNO), ""),
    ("7c", "Data Governance & Security",
     "Must all connections travel over a private internal network, not the public internet? (AWS PrivateLink, Azure Private Link), if yes, estimated [GB processed/month].",
     "PrivateLink adds a per-endpoint monthly charge plus a per-GB data-processed charge, and requires Enterprise Edition or higher.",
     ("list+", YESNO), "if yes: [GB processed/month]"),

    # 8, Disaster Recovery & Business Continuity
    ("8a", "Disaster Recovery & Business Continuity",
     "Do you require a geographically separate backup/failover copy in case the primary region is unavailable? (yes / no)",
     "Cross-region replication is a frequently overlooked cost (seed, ongoing replication compute ~4 cr/TB, egress, replica storage)",
     ("list", YESNO), ""),
    ("8b", "Disaster Recovery & Business Continuity",
     "[If yes] How often should the backup sync? And what is your production dataset at go-live [TB] and daily change volume [GB/day]?",
     "Sync frequency sets the compute cadence; dataset size drives the one-time seed transfer; daily change drives ongoing monthly compute and egress.",
     "text", "sync freq, [TB], [GB/day]"),
    ("8c", "Disaster Recovery & Business Continuity",
     '[If yes] Which regions are involved? (source -> target, e.g. "AWS US East -> AWS EU West")',
     "Egress varies by region pair. The pair also determines whether an Egress Cost Optimizer cache applies.",
     "text", "source -> target"),

    # 9, Real-time Database Replication (CDC)
    ("9a", "Real-time Database Replication (CDC)",
     "Do you need to continuously capture and replicate row-level changes from operational databases in near real-time? (yes / no)",
     "CDC is a separate cost category: a dedicated runtime service plus a warehouse for MERGE operations.",
     ("list", YESNO), ""),
    ("9b", "Real-time Database Replication (CDC)",
     "[If yes] Which source database types? How many separate database servers? Same physical server? For Oracle: processor type (x86 / SPARC / POWER) and [# cores].",
     "Each distinct source server needs a dedicated runtime node. Oracle adds per-core licensing; the processor type sets the licensing factor (x86 = 0.5, SPARC/POWER = 1.0).",
     "text", "DB types, # servers, Oracle cores"),
    ("9c", "Real-time Database Replication (CDC)",
     "[If yes] Roughly how many [# tables] need replicating across all sources, and approximate [row changes per day]?",
     "Runtime nodes have a 600-table limit. Daily event volume sets runtime size: < 100k events/day -> Medium; > 100k -> Large.",
     "text", "[# tables] + [row changes/day]"),
    ("9d", "Real-time Database Replication (CDC)",
     "[If yes] Is there an existing warehouse that can be reused for the CDC MERGE step, or must a new one be provisioned? (reuse existing / provision new)",
     "The MERGE warehouse is typically 60-70% of total CDC cost. Reusing an existing warehouse avoids duplicating that cost in the estimate.",
     ("list+", ["Reuse existing", "Provision new"]), ""),

    # 10, Data Sharing & External Collaboration
    ("10a", "Data Sharing & External Collaboration",
     "Do you need to share live data with external partners, customers, or vendors outside your organisation?, if yes: [# external parties] and how often they query.",
     "External sharing to parties without their own account needs a Reader Account (billed to you). A dedicated warehouse must be sized from external user count and query frequency.",
     "text", "if yes: [# parties] + query freq"),
    ("10b", "Data Sharing & External Collaboration",
     "Are you planning to sell data products/apps via a marketplace, or build a data-powered app your customers will use?",
     "Determines whether Native Apps / Marketplace costs apply; these also require Enterprise Edition or higher.",
     "text", ""),

    # 11, Future Growth & Platform
    ("11a", "Future Growth & Platform",
     "Projected annual growth in data volume over the next 3 years? [% per year]",
     "Essential for 3-year TCO. Storage is a recurring cost that scales with data growth.",
     "pct", "[% per year]"),
    ("11b", "Future Growth & Platform",
     "Projected annual growth in number of users and/or new workloads over the next 3 years? [% per year]",
     "Essential for modeling rising compute consumption across the X-year term, ensuring the contract value is accurate for the full duration.",
     "pct", "[% per year]"),
]


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def plugin_version():
    """Read the latest version number from CHANGELOG.md (first '## [x.y.z]')."""
    try:
        with open(CHANGELOG_PATH, encoding="utf-8") as fh:
            for line in fh:
                m = re.match(r"##\s*\[([^\]]+)\]", line.strip())
                if m:
                    return m.group(1)
    except OSError:
        pass
    return "unknown"


def _set_cell_background(cell, hex_color):
    """Shade a table cell (python-docx has no built-in cell-fill API)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, color=BORDER_GREY, sz="8"):
    """Apply thin borders on all four sides of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _style_run(run, *, bold=False, italic=False, color="222222", size=BODY_PT):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_text(cell, text, *, bold=False, color="222222", size=BODY_PT,
                   align=WD_ALIGN_PARAGRAPH.LEFT):
    """Replace a cell's content with a single styled paragraph/run."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    _style_run(run, bold=bold, color=color, size=size)
    return para


def _set_repeat_header(row):
    """Mark a table row to repeat as a header on each page."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _answer_guidance(atype, hint):
    """Return guidance text to display inside an empty Answer cell, or ''. """
    parts = []
    if isinstance(atype, tuple):
        _, options = atype
        parts.append("Options: " + " / ".join(options))
    if hint:
        parts.append(hint)
    return "   ".join(parts)


def _input_prompt(atype, hint):
    """Return placeholder text for non-list answer types.

    Mirrors the Excel _make_validation prompt text so both outputs give
    identical guidance to the respondent.
    """
    if atype == "num":
        return f"Numeric value{(' ' + hint) if hint else ''}. Leave blank if unknown."
    if atype == "pct":
        return f"A percentage 0–100{(' ' + hint) if hint else ''}. Leave blank if unknown."
    return hint  # "text" + unit/guidance hint, or "" for plain text with no hint


# Monotonic id source for content controls (each w:sdt needs a unique w:id).
_SDT_ID = [90000]


def _add_dropdown(cell, options, *, combo, tag, placeholder="Choose an option…"):
    """Insert a Word dropdown content control into a cell's first paragraph.

    This is the Word equivalent of Excel data validation: ``combo=False`` is a
    strict drop-down list (only listed values), ``combo=True`` is a combo box
    (listed values, but the user may also type their own).
    """
    para = cell.paragraphs[0]

    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(BODY_PT * 2))  # half-points
    rPr.append(sz)
    sdtPr.append(rPr)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "Answer")
    sdtPr.append(alias)

    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdtPr.append(tag_el)

    _SDT_ID[0] += 1
    id_el = OxmlElement("w:id")
    id_el.set(qn("w:val"), str(_SDT_ID[0]))
    sdtPr.append(id_el)

    list_el = OxmlElement("w:comboBox" if combo else "w:dropDownList")
    placeholder_item = OxmlElement("w:listItem")
    placeholder_item.set(qn("w:displayText"), placeholder)
    placeholder_item.set(qn("w:value"), "")
    list_el.append(placeholder_item)
    for opt in options:
        li = OxmlElement("w:listItem")
        li.set(qn("w:displayText"), opt)
        li.set(qn("w:value"), opt)
        list_el.append(li)
    sdtPr.append(list_el)

    sdt.append(sdtPr)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    run_rpr = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), FONT)
    run_fonts.set(qn("w:hAnsi"), FONT)
    run_rpr.append(run_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREY)
    run_rpr.append(color)
    run_sz = OxmlElement("w:sz")
    run_sz.set(qn("w:val"), str(BODY_PT * 2))
    run_rpr.append(run_sz)
    run.append(run_rpr)
    t = OxmlElement("w:t")
    t.text = placeholder
    run.append(t)
    sdt_content.append(run)
    sdt.append(sdt_content)

    para._p.append(sdt)


def _add_plain_text_control(cell, placeholder, *, tag):
    """Insert an inline plain-text content control with grey italic placeholder.

    Word equivalent of an Excel DataValidation input-message for numeric and
    free-text answer types (num / pct / text+hint).  When the respondent clicks
    the cell and types, the placeholder is replaced by their answer.
    """
    para = cell.paragraphs[0]

    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(BODY_PT * 2))
    rPr.append(sz)
    sdtPr.append(rPr)

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "Answer")
    sdtPr.append(alias)

    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdtPr.append(tag_el)

    _SDT_ID[0] += 1
    id_el = OxmlElement("w:id")
    id_el.set(qn("w:val"), str(_SDT_ID[0]))
    sdtPr.append(id_el)

    # w:text marks this as a plain-text (single-value) content control
    text_el = OxmlElement("w:text")
    sdtPr.append(text_el)

    sdt.append(sdtPr)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    run_rpr = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), FONT)
    run_fonts.set(qn("w:hAnsi"), FONT)
    run_rpr.append(run_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), GREY)
    run_rpr.append(color_el)
    run_sz = OxmlElement("w:sz")
    run_sz.set(qn("w:val"), str((BODY_PT - 3) * 2))
    run_rpr.append(run_sz)
    run_rpr.append(OxmlElement("w:i"))
    run.append(run_rpr)
    t = OxmlElement("w:t")
    t.text = placeholder
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    sdt_content.append(run)
    sdt.append(sdt_content)

    para._p.append(sdt)


def _set_col_widths(table, widths_in):
    """Force fixed column widths (set on every cell for Word to honour them)."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_in):
            row.cells[idx].width = Inches(w)


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build():
    doc = Document()

    # Landscape, slim margins -> enough usable width for the 4-column table
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # default document font
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)

    # --------------------------------------------------------------- Cover page
    if os.path.exists(LOGO_PATH):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_para.add_run().add_picture(LOGO_PATH, height=Inches(0.9))

    title = doc.add_paragraph()
    title_run = title.add_run("Snowflake Solution Sizing Assessment")
    _style_run(title_run, bold=True, color=DARK_BLUE, size=26)
    title.space_after = Pt(12)

    intro = doc.add_paragraph()
    _style_run(
        intro.add_run(
            "The details you provide enable us to build an accurate, right-sized "
            "solution and a predictable cost estimate tailored to your goals."
        ),
        color="333333",
    )

    how = doc.add_paragraph()
    _style_run(how.add_run("How to answer"), bold=True, color=BLUE, size=BODY_PT + 1)

    bullets = [
        "Where a question shows a unit in square brackets (e.g. [# concurrent users], "
        "[GB/day], [indexed GB]), please give a number in that unit. A rough number is "
        "far better than none.",
        "Leave an answer blank if genuinely unknown; anything blank is sized with a "
        "documented assumption and flagged for confirmation.",
        "Enter your answers in the highlighted Answer column in the table that follows. "
        "Where options are listed, choose one of them.",
    ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        _style_run(p.add_run(b), color="333333")

    tip = doc.add_paragraph()
    _style_run(
        tip.add_run("Tip: blue-shaded cells are for you to fill in."),
        italic=True, color=DARK_BLUE,
    )

    ver = doc.add_paragraph()
    _style_run(ver.add_run(f"version {plugin_version()}"), color=GREY, size=BODY_PT - 2)

    doc.add_page_break()

    # ------------------------------------------------------- Questionnaire table
    headers = ["ID", "Question", "Answer", "Comments"]
    col_widths = [0.5, 4.3, 2.8, 2.2]  # inches (total 9.8" = landscape usable width)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    hdr = table.rows[0]
    _set_repeat_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        _set_cell_background(cell, DARK_BLUE)
        _set_cell_borders(cell)
        _set_cell_text(cell, h, bold=True, color=WHITE)

    current_cat = None
    band = False  # alternating tint within a category

    for (qid, cat, question, reason, atype, hint) in QUESTIONS:
        # category divider row
        if cat != current_cat:
            current_cat = cat
            band = False
            crow = table.add_row()
            a = crow.cells[0]
            merged = a.merge(crow.cells[len(headers) - 1])
            _set_cell_background(merged, BLUE)
            _set_cell_borders(merged)
            _set_cell_text(merged, cat, bold=True, color=WHITE)

        tint = ROW_TINT if band else WHITE
        band = not band

        r = table.add_row()
        c_id, c_q, c_a, c_com = r.cells

        _set_cell_text(c_id, qid, bold=True, color=DARK_BLUE,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(c_q, question, color="222222")
        for cell in (c_id, c_q):
            _set_cell_background(cell, tint)

        # Answer cell: empty + shaded.
        # Validation mirrors the Excel _make_validation logic exactly:
        #   list  → strict Word drop-down list  (no other values)
        #   list+ → Word combo box             (listed values + free entry)
        #   num   → plain-text control, prompt "Numeric value … Leave blank if unknown."
        #   pct   → plain-text control, prompt "A percentage 0–100 … Leave blank if unknown."
        #   text  → plain-text control showing the unit/guidance hint (if any)
        #   text  (no hint) → plain shaded cell, no control (matches Excel returning None)
        c_a.text = ""
        _set_cell_background(c_a, ANSWER_FILL)
        if isinstance(atype, tuple):
            kind, options = atype
            _add_dropdown(c_a, options, combo=(kind == "list+"), tag=f"q_{qid}")
            if hint:
                gp = c_a.add_paragraph()
                _style_run(gp.add_run(hint), italic=True, color=GREY,
                           size=BODY_PT - 3)
        else:
            prompt = _input_prompt(atype, hint)
            if prompt:
                _add_plain_text_control(c_a, prompt, tag=f"q_{qid}")

        # Comments cell: empty + shaded
        c_com.text = ""
        _set_cell_background(c_com, ANSWER_FILL)

        for cell in (c_id, c_q, c_a, c_com):
            _set_cell_borders(cell)

    _set_col_widths(table, col_widths)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
